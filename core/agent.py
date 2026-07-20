"""
Pleadly Agent 工具调用模块
负责: WebSearch联网检索 / 文档解析 / 工作流编排
"""

import os
import json
import requests
from typing import Optional, List, Dict, Any
from dataclasses import dataclass

# ── Web Search Tool ──

def web_search(query: str, max_results: int = 5) -> List[Dict[str, str]]:
    """
    联网检索工具 — 使用 DuckDuckGo 免费搜索(无需API Key)
    备选: 可切换为 Serper.dev 或 Bing Search API
    """
    results = []
    try:
        from duckduckgo_search import DDGS
        with DDGS() as ddgs:
            for r in ddgs.text(query, max_results=max_results):
                results.append({
                    "title": r.get("title", ""),
                    "url": r.get("href", ""),
                    "snippet": r.get("body", "")[:300],
                })
    except ImportError:
        # Fallback: basic requests + BeautifulSoup approach
        results.append({
            "title": "搜索功能需安装 duckduckgo-search 包",
            "url": "",
            "snippet": "请运行: pip install duckduckgo-search"
        })
    return results

def search_company_info(company_name: str) -> str:
    """专门搜索公司背景信息，返回格式化文本。"""
    queries = [
        f"{company_name} 公司介绍 主营业务 规模",
        f"{company_name} 招聘 薪资 评价",
        f"{company_name} AI 数字化 技术",
    ]
    all_info = []
    for q in queries:
        results = web_search(q, max_results=3)
        for r in results:
            all_info.append(f"【{r['title']}】\n{r['snippet']}\n来源: {r['url']}")
    return "\n\n".join(all_info[:8])  # Cap at 8 snippets

def search_interview_experience(company: str, role: str) -> str:
    """搜索真实面试经验。"""
    queries = [
        f"{company} {role} 面试 面经 2025 2026",
        f"{company} 校招 面试题",
        f"{role} 面试 常见问题 应届生",
    ]
    all_info = []
    for q in queries:
        results = web_search(q, max_results=3)
        for r in results:
            all_info.append(f"【{r['title']}】\n{r['snippet']}\n来源: {r['url']}")
    return "\n\n".join(all_info[:10])

# ── Document Parser ──

def parse_docx(file_path: str) -> str:
    """解析Word文档。"""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Also parse tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        return f"[文档解析失败: {e}]"

def parse_pdf(file_path: str) -> str:
    """解析PDF文档。"""
    try:
        from pypdf2 import PdfReader
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    except Exception as e:
        return f"[PDF解析失败: {e}]"

def parse_resume_file(file_path: str) -> str:
    """自动检测文件类型并解析简历。"""
    if not os.path.exists(file_path):
        return f"[文件不存在: {file_path}]"

    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.docx':
        return parse_docx(file_path)
    elif ext == '.pdf':
        return parse_pdf(file_path)
    elif ext in ['.txt', '.md']:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    else:
        return f"[不支持的文件格式: {ext}]"

# ── Salary Data Tool ──

def search_salary_data(role: str, city: str = "深圳") -> str:
    """搜索岗位薪资数据。"""
    queries = [
        f"{city} {role} 薪资 2026 应届生 校招",
        f"{role} 薪资范围 中位数 {city}",
    ]
    all_info = []
    for q in queries:
        results = web_search(q, max_results=3)
        for r in results:
            all_info.append(f"【{r['title']}】\n{r['snippet']}")
    return "\n\n".join(all_info[:6])

# ── Agent Orchestrator ──

@dataclass
class AgentState:
    resume_text: str = ""
    jd_text: str = ""
    company_research: str = ""
    salary_research: str = ""
    interview_experiences: str = ""
    step_results: Dict[str, str] = None

    def __post_init__(self):
        if self.step_results is None:
            self.step_results = {}

class AgentOrchestrator:
    """Agent 编排器 — 管理工作流执行顺序和工具调用。"""

    def __init__(self):
        self.state = AgentState()

    def set_input(self, resume: str, jd: str):
        self.state.resume_text = resume
        self.state.jd_text = jd

    def research(self, company_name: str = "") -> str:
        """执行预研究：公司背景 + 薪资 + 面经。"""
        if company_name:
            self.state.company_research = search_company_info(company_name)

        # Try to extract role from JD
        role_hint = "AI产品经理"  # Default, will be overridden
        if self.state.jd_text:
            # Extract first line or title from JD
            lines = self.state.jd_text.strip().split('\n')
            if lines:
                role_hint = lines[0][:30]

        self.state.salary_research = search_salary_data(role_hint)

        return f"预研究完成。公司信息: {len(self.state.company_research)}字符。薪资数据: {len(self.state.salary_research)}字符。"

    def get_rag_context(self) -> str:
        """组装RAG检索上下文。"""
        parts = []
        if self.state.company_research:
            parts.append(f"## 公司研究\n{self.state.company_research[:3000]}")
        if self.state.salary_research:
            parts.append(f"## 薪资参考\n{self.state.salary_research[:1500]}")
        if self.state.interview_experiences:
            parts.append(f"## 面经参考\n{self.state.interview_experiences[:2000]}")
        return "\n\n".join(parts)
